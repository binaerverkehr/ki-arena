"""
Dokument-Verarbeitung – Extraktion von Text und Aufbereitung von Bildern.

Unterstützte Formate:
- Text: PDF, DOCX, TXT, MD
- Bilder: PNG, JPG, JPEG, WEBP

Extrahierter Text wird als Referenzmaterial in Systemprompts injiziert.
Bilder werden als Base64 für Vision-APIs aufbereitet.
"""
from __future__ import annotations

import base64
import io
import logging
from dataclasses import dataclass
from pathlib import Path

from fastapi import UploadFile

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
TEXT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_FILES_PER_ROLE = 5
MAX_TEXT_CHARS = 15_000  # Pro Dokument

MIME_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
}


@dataclass
class ProcessedDocument:
    filename: str
    file_type: str  # "text" | "image"
    extracted_text: str = ""
    base64_data: str = ""
    media_type: str = ""
    error: str = ""


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_pdf(data: bytes, lang: str = "de") -> str:
    """Extrahiert Text aus einer PDF-Datei mit pdfplumber."""
    import pdfplumber

    page_label = "Seite" if lang == "de" else "Page"
    parts: list[str] = []

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                parts.append(f"[{page_label} {i}]\n{text}")

    full_text = "\n\n".join(parts)
    if len(full_text) > MAX_TEXT_CHARS:
        full_text = full_text[:MAX_TEXT_CHARS] + "\n[… gekürzt]"
    return full_text


def _extract_docx(data: bytes) -> str:
    """Extrahiert Text aus einer DOCX-Datei."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tabellen ebenfalls extrahieren
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    full_text = "\n\n".join(parts)
    if len(full_text) > MAX_TEXT_CHARS:
        full_text = full_text[:MAX_TEXT_CHARS] + "\n[… gekürzt]"
    return full_text


def _extract_text_file(data: bytes) -> str:
    """Liest eine TXT/MD-Datei."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")

    text = text.strip()
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n[… gekürzt]"
    return text


# ---------------------------------------------------------------------------
# Image processing
# ---------------------------------------------------------------------------

def _process_image(data: bytes, ext: str) -> tuple[str, str]:
    """Validiert und bereitet ein Bild auf. Gibt (base64_data, media_type) zurück."""
    from PIL import Image

    img = Image.open(io.BytesIO(data))
    img.verify()  # Prüft ob das Bild valide ist

    # Nochmal öffnen nach verify (verify macht das Objekt unbrauchbar)
    img = Image.open(io.BytesIO(data))

    # Resize wenn nötig (max 2048px auf längster Seite)
    max_dim = 2048
    if max(img.size) > max_dim:
        ratio = max_dim / max(img.size)
        new_size = (int(img.width * ratio), int(img.height * ratio))
        img = img.resize(new_size, Image.Resampling.LANCZOS)

    # In das passende Format konvertieren
    output = io.BytesIO()
    media_type = MIME_TYPES.get(ext, "image/png")

    if ext == ".webp":
        img.save(output, format="WEBP")
    elif ext in (".jpg", ".jpeg"):
        if img.mode == "RGBA":
            img = img.convert("RGB")
        img.save(output, format="JPEG", quality=85)
    else:
        img.save(output, format="PNG")

    b64 = base64.b64encode(output.getvalue()).decode("ascii")
    return b64, media_type


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def process_upload(file: UploadFile, lang: str = "de") -> ProcessedDocument:
    """Verarbeitet eine einzelne hochgeladene Datei."""
    filename = file.filename or "unknown"
    ext = Path(filename).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        return ProcessedDocument(
            filename=filename, file_type="text",
            error=f"Nicht unterstütztes Format: {ext}",
        )

    data = await file.read()

    if len(data) > MAX_FILE_SIZE:
        return ProcessedDocument(
            filename=filename, file_type="text",
            error=f"Datei zu groß: {len(data) / 1024 / 1024:.1f} MB (max. {MAX_FILE_SIZE // 1024 // 1024} MB)",
        )

    try:
        if ext in IMAGE_EXTENSIONS:
            b64, media_type = _process_image(data, ext)
            return ProcessedDocument(
                filename=filename, file_type="image",
                base64_data=b64, media_type=media_type,
            )
        elif ext == ".pdf":
            text = _extract_pdf(data, lang)
            return ProcessedDocument(filename=filename, file_type="text", extracted_text=text)
        elif ext == ".docx":
            text = _extract_docx(data)
            return ProcessedDocument(filename=filename, file_type="text", extracted_text=text)
        else:  # .txt, .md
            text = _extract_text_file(data)
            return ProcessedDocument(filename=filename, file_type="text", extracted_text=text)
    except Exception as e:
        logger.exception("Fehler bei Verarbeitung von %s", filename)
        return ProcessedDocument(
            filename=filename, file_type="text",
            error=f"Verarbeitungsfehler: {e}",
        )


async def process_uploads(files: list[UploadFile], lang: str = "de") -> list[ProcessedDocument]:
    """Verarbeitet eine Liste von Uploads."""
    results: list[ProcessedDocument] = []
    for file in files[:MAX_FILES_PER_ROLE]:
        doc = await process_upload(file, lang)
        results.append(doc)
    return results


def build_context_block(docs: list[ProcessedDocument], lang: str = "de") -> str:
    """Formatiert extrahierte Texte als Referenzmaterial-Block für den Systemprompt."""
    text_docs = [d for d in docs if d.file_type == "text" and d.extracted_text and not d.error]
    if not text_docs:
        return ""

    if lang == "de":
        header = "--- Referenzmaterial ---\nNutze die folgenden Informationen, um deine Argumente mit konkreten Fakten und Zitaten zu untermauern:\n"
    else:
        header = "--- Reference Materials ---\nUse the following information to support your arguments with specific facts and citations:\n"

    parts = [header]
    for doc in text_docs:
        parts.append(f"\n[Datei: {doc.filename}]\n{doc.extracted_text}")

    parts.append("\n---")
    return "\n".join(parts)


def get_image_attachments(docs: list[ProcessedDocument]) -> list[dict]:
    """Extrahiert Bild-Attachments als Liste von Dicts für die LLM-Provider."""
    return [
        {
            "filename": d.filename,
            "base64_data": d.base64_data,
            "media_type": d.media_type,
        }
        for d in docs
        if d.file_type == "image" and d.base64_data and not d.error
    ]
